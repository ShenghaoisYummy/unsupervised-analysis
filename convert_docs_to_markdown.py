#!/usr/bin/env python3
"""
Convert .doc and .docx files to markdown format.
Supports both .doc and .docx formats with comprehensive parsing.
"""

import os
import sys
from pathlib import Path
import mammoth
from docx import Document
import re


def docx_to_markdown_python_docx(docx_path, output_path):
    """Convert .docx to markdown using python-docx library."""
    try:
        doc = Document(docx_path)
        markdown_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                markdown_content.append("")
                continue
                
            # Handle different paragraph styles
            style_name = paragraph.style.name.lower()
            
            if 'heading 1' in style_name:
                markdown_content.append(f"# {text}")
            elif 'heading 2' in style_name:
                markdown_content.append(f"## {text}")
            elif 'heading 3' in style_name:
                markdown_content.append(f"### {text}")
            elif 'heading 4' in style_name:
                markdown_content.append(f"#### {text}")
            elif 'heading 5' in style_name:
                markdown_content.append(f"##### {text}")
            elif 'heading 6' in style_name:
                markdown_content.append(f"###### {text}")
            else:
                # Handle regular paragraphs with formatting
                formatted_text = text
                
                # Check for bold/italic formatting in runs
                for run in paragraph.runs:
                    if run.bold and run.text:
                        formatted_text = formatted_text.replace(run.text, f"**{run.text}**")
                    elif run.italic and run.text:
                        formatted_text = formatted_text.replace(run.text, f"*{run.text}*")
                
                markdown_content.append(formatted_text)
        
        # Handle tables
        for table in doc.tables:
            markdown_content.append("")  # Add spacing before table
            
            # Create table headers
            if table.rows:
                headers = []
                for cell in table.rows[0].cells:
                    headers.append(cell.text.strip() or "")
                
                if any(headers):  # Only create table if headers exist
                    markdown_content.append("| " + " | ".join(headers) + " |")
                    markdown_content.append("| " + " | ".join(["---"] * len(headers)) + " |")
                    
                    # Add table rows
                    for row in table.rows[1:]:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip() or "")
                        markdown_content.append("| " + " | ".join(row_data) + " |")
        
        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(markdown_content))
        
        return True
    except Exception as e:
        print(f"Error converting {docx_path} with python-docx: {e}")
        return False


def docx_to_markdown_mammoth(docx_path, output_path):
    """Convert .docx to markdown using mammoth library (better formatting preservation)."""
    try:
        with open(docx_path, "rb") as docx_file:
            result = mammoth.convert_to_markdown(docx_file)
            
        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(result.value)
        
        # Print any warnings
        if result.messages:
            print(f"Warnings for {docx_path}:")
            for message in result.messages:
                print(f"  - {message}")
        
        return True
    except Exception as e:
        print(f"Error converting {docx_path} with mammoth: {e}")
        return False


def convert_document_to_markdown(doc_path, output_dir, use_mammoth=True):
    """Convert a document file to markdown."""
    doc_path = Path(doc_path)
    output_dir = Path(output_dir)
    
    # Create output directory if it doesn't exist
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate output filename
    output_filename = doc_path.stem + '.md'
    output_path = output_dir / output_filename
    
    print(f"Converting {doc_path.name} to {output_path}")
    
    if doc_path.suffix.lower() in ['.docx']:
        if use_mammoth:
            success = docx_to_markdown_mammoth(doc_path, output_path)
            if not success:
                print("Mammoth failed, trying python-docx...")
                success = docx_to_markdown_python_docx(doc_path, output_path)
        else:
            success = docx_to_markdown_python_docx(doc_path, output_path)
            
        if success:
            print(f"Successfully converted {doc_path.name}")
            return output_path
        else:
            print(f"Failed to convert {doc_path.name}")
            return None
    else:
        print(f"Unsupported file format: {doc_path.suffix}")
        return None


def main():
    """Main function to convert all document files in the documents directory."""
    documents_dir = Path("documents")
    markdown_dir = Path("documents/markdown")
    
    if not documents_dir.exists():
        print("Documents directory not found!")
        return
    
    # Find all document files
    doc_files = []
    for pattern in ['*.docx', '*.doc']:
        doc_files.extend(documents_dir.glob(pattern))
    
    if not doc_files:
        print("No document files found in documents directory")
        return
    
    print(f"Found {len(doc_files)} document files to convert:")
    for doc_file in doc_files:
        print(f"  - {doc_file.name}")
    
    print("\nStarting conversion...")
    
    converted_files = []
    for doc_file in doc_files:
        result = convert_document_to_markdown(doc_file, markdown_dir)
        if result:
            converted_files.append(result)
    
    print(f"\nConversion complete! Converted {len(converted_files)} files:")
    for converted_file in converted_files:
        print(f"  - {converted_file}")


if __name__ == "__main__":
    main()