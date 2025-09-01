import os
import re
from pathlib import Path
from docx import Document
from typing import List, Optional
from config import Config

class DocxToMarkdownParser:
    def __init__(self):
        self.config = Config()
        
    def clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        text = text.strip()
        return text
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text content from Word document"""
        try:
            doc = Document(docx_path)
            content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Handle different paragraph styles
                    style = paragraph.style.name if paragraph.style else "Normal"
                    
                    if "Heading" in style:
                        level = 1
                        if "Heading 1" in style:
                            level = 1
                        elif "Heading 2" in style:
                            level = 2
                        elif "Heading 3" in style:
                            level = 3
                        elif "Heading 4" in style:
                            level = 4
                        else:
                            level = min(int(re.search(r'\d+', style).group()) if re.search(r'\d+', style) else 1, 6)
                        
                        content.append(f"{'#' * level} {text}")
                    else:
                        content.append(text)
            
            # Handle tables
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_content.append(cell_text)
                    if any(row_content):  # Only add non-empty rows
                        table_content.append("| " + " | ".join(row_content) + " |")
                
                if table_content:
                    # Add table header separator for first row
                    if len(table_content) > 0:
                        header_sep = "| " + " | ".join(["---"] * len(table_content[0].split("|")[1:-1])) + " |"
                        table_content.insert(1, header_sep)
                    
                    content.append("\n" + "\n".join(table_content) + "\n")
            
            return self.clean_text("\n\n".join(content))
        
        except Exception as e:
            raise Exception(f"Error extracting text from {docx_path}: {str(e)}")
    
    def save_as_markdown(self, content: str, output_path: str, title: Optional[str] = None):
        """Save content as markdown file"""
        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                if title:
                    f.write(f"# {title}\n\n")
                f.write(content)
                
        except Exception as e:
            raise Exception(f"Error saving markdown to {output_path}: {str(e)}")
    
    def convert_single_document(self, docx_path: str, output_path: Optional[str] = None) -> str:
        """Convert a single Word document to markdown"""
        if not output_path:
            docx_name = Path(docx_path).stem
            output_path = os.path.join(self.config.INTERVIEWS_DIR, f"{docx_name}.md")
        
        # Extract content
        content = self.extract_text_from_docx(docx_path)
        
        # Use filename as title
        title = Path(docx_path).stem.replace('_', ' ').replace('-', ' ').title()
        
        # Save as markdown
        self.save_as_markdown(content, output_path, title)
        
        return output_path
    
    def convert_batch(self, documents_dir: Optional[str] = None) -> List[str]:
        """Convert all Word documents in a directory to markdown"""
        if not documents_dir:
            documents_dir = self.config.DOCUMENTS_DIR
        
        if not os.path.exists(documents_dir):
            raise FileNotFoundError(f"Documents directory not found: {documents_dir}")
        
        # Find all docx files
        docx_files = []
        for ext in ['*.docx', '*.doc']:
            docx_files.extend(Path(documents_dir).glob(ext))
        
        if not docx_files:
            raise FileNotFoundError(f"No Word documents found in {documents_dir}")
        
        converted_files = []
        failed_conversions = []
        
        print(f"Found {len(docx_files)} Word documents to convert...")
        
        for docx_file in docx_files:
            try:
                output_path = self.convert_single_document(str(docx_file))
                converted_files.append(output_path)
                print(f"✓ Converted: {docx_file.name} → {Path(output_path).name}")
            except Exception as e:
                failed_conversions.append((str(docx_file), str(e)))
                print(f"✗ Failed: {docx_file.name} - {str(e)}")
        
        if failed_conversions:
            print(f"\nConversion Summary:")
            print(f"✓ Successful: {len(converted_files)}")
            print(f"✗ Failed: {len(failed_conversions)}")
            for failed_file, error in failed_conversions:
                print(f"  - {Path(failed_file).name}: {error}")
        
        return converted_files
    
    def get_converted_files(self) -> List[str]:
        """Get list of all converted markdown files"""
        interviews_dir = Path(self.config.INTERVIEWS_DIR)
        if not interviews_dir.exists():
            return []
        
        return [str(f) for f in interviews_dir.glob("*.md")]

if __name__ == "__main__":
    parser = DocxToMarkdownParser()
    
    # Convert all documents
    try:
        converted_files = parser.convert_batch()
        print(f"\nSuccessfully converted {len(converted_files)} documents to markdown format.")
        print("Files saved in:", Config.INTERVIEWS_DIR)
    except Exception as e:
        print(f"Error during batch conversion: {e}")